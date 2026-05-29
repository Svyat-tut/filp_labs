from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from datetime import datetime
import shutil


class LetterGenerator:
    def __init__(self, template_path="template.docx"):
        self.template_path = template_path
        self.output_dir = "generated_letters"
        

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generate(self, replacements):


        doc = Document(self.template_path)
        

        self._replace_in_paragraphs(doc, replacements)
        

        self._replace_in_tables(doc, replacements)
        

        self._replace_in_headers_footers(doc, replacements)
        

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"letter_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        

        doc.save(output_path)
        
        return output_path
    
    def _replace_in_paragraphs(self, doc, replacements):

        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    self._replace_text_in_paragraph(paragraph, placeholder, value)
    
    def _replace_in_tables(self, doc, replacements):

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, placeholder, value)
    
    def _replace_in_headers_footers(self, doc, replacements):

        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self._replace_text_in_paragraph(paragraph, placeholder, value)
            

            footer = section.footer
            for paragraph in footer.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self._replace_text_in_paragraph(paragraph, placeholder, value)
    
    def _replace_text_in_paragraph(self, paragraph, placeholder, replacement_value):

        if placeholder not in paragraph.text:
            return
        

        for i, run in enumerate(paragraph.runs):
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement_value)
                return
        

        full_text = "".join([run.text for run in paragraph.runs])
        
        if placeholder in full_text:

            new_text = full_text.replace(placeholder, replacement_value)
            

            for run in paragraph.runs:
                run.text = ""
            

            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.text = new_text
