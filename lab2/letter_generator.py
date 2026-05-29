from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


class Attachment:

    def __init__(self, title, content, page_count=1):

        self.title = title
        self.content = content
        self.page_count = page_count


class LetterGenerator:
    def __init__(self, template_path="template.docx"):
        self.template_path = template_path
        self.output_dir = "generated_letters"
        

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generate(self, replacements, attachments=None):

        if attachments is None:
            attachments = []
        

        doc = Document(self.template_path)
        

        self._replace_in_paragraphs(doc, replacements)
        

        self._replace_in_tables(doc, replacements)
        

        self._replace_in_headers_footers(doc, replacements)
        

        if attachments:
            self._add_attachments_section(doc, attachments)
        

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"letter_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        

        doc.save(output_path)
        
        return output_path
    
    def _replace_in_paragraphs(self, doc, replacements):

        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    self._replace_text_in_paragraph(paragraph, placeholder, value)
    
    def _replace_in_tables(self, doc, replacements):

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, placeholder, value)
    
    def _replace_in_headers_footers(self, doc, replacements):

        for section in doc.sections:

            header = section.header
            for paragraph in header.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self._replace_text_in_paragraph(paragraph, placeholder, value)
            

            footer = section.footer
            for paragraph in footer.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self._replace_text_in_paragraph(paragraph, placeholder, value)
    
    def _replace_text_in_paragraph(self, paragraph, placeholder, replacement_value):

        if placeholder not in paragraph.text:
            return
        

        for i, run in enumerate(paragraph.runs):
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement_value)
                return
        

        full_text = "".join([run.text for run in paragraph.runs])
        
        if placeholder in full_text:

            new_text = full_text.replace(placeholder, replacement_value)
            

            for run in paragraph.runs:
                run.text = ""
            

            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.text = new_text
    
    def _add_attachments_section(self, doc, attachments):

        doc.add_paragraph()
        

        attachments_header = doc.add_paragraph()
        attachments_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
        attachments_header_run = attachments_header.add_run("Приложения:")
        attachments_header_run.font.bold = True
        

        for i, attachment in enumerate(attachments, 1):
            attachment_para = doc.add_paragraph(
                f"{attachment.title} на {attachment.page_count} л."
            )
            attachment_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_page_break()
        

        for i, attachment in enumerate(attachments, 1):
            attachment_header = doc.add_paragraph()
            attachment_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if len(attachments) == 1:
                header_text = "Приложение"
            else:
                header_text = f"Приложение {i}"
            
            header_run = attachment_header.add_run(header_text)
            header_run.font.bold = True
            
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(attachment.title)
            title_run.font.bold = True
            title_run.font.size = Pt(12)
            
            doc.add_paragraph()
            
            content_lines = attachment.content.split('\n')
            for line in content_lines:
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.add_paragraph()
            

            if i < len(attachments):
                doc.add_page_break()

