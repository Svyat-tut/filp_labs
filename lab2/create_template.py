from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():

    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ПИСЬМО")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    

    doc.add_paragraph()
    date_para = doc.add_paragraph("Дата: 25 мая 2026 г.")
    

    doc.add_paragraph()
    doc.add_paragraph("Кому:")
    recipient = doc.add_paragraph("{RECIPIENT_POST}")
    recipient.paragraph_format.left_indent = Inches(0.5)
    recipient = doc.add_paragraph("{RECIPIENT_NAME}")
    recipient.paragraph_format.left_indent = Inches(0.5)
    recipient = doc.add_paragraph("{RECIPIENT_ORG}")
    recipient.paragraph_format.left_indent = Inches(0.5)
    

    doc.add_paragraph()
    doc.add_paragraph("Тема: {LETTER_SUBJECT}")
    

    doc.add_paragraph()
    doc.add_paragraph("{LETTER_BODY}")
    

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("С уважением,")
    doc.add_paragraph()
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Директор")
    

    doc.save("template.docx")
    print("✓ Шаблон 'template.docx' успешно создан!")


if __name__ == "__main__":
    create_template()
