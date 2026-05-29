import tkinter as tk
from tkinter import ttk, messagebox
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping
import tempfile
from PIL import Image, ImageTk
import fitz


def main():
    root = tk.Tk()
    root.title("Генератор документов с приложениями")
    root.geometry("1440x860")
    root.configure(bg="#f0f2f5")
    root.minsize(1100, 600)

    # --- Цветовая палитра ---
    CLR_BG        = "#f0f2f5"
    CLR_PANEL     = "#ffffff"
    CLR_ACCENT    = "#4f6df5"
    CLR_ACCENT_H  = "#3a57d4"
    CLR_DANGER    = "#e84545"
    CLR_DANGER_H  = "#c73232"
    CLR_TEXT      = "#1e2235"
    CLR_SUBTEXT   = "#6b7280"
    CLR_BORDER    = "#e2e5ec"
    CLR_ENTRY_BG  = "#f8f9fc"

    style = ttk.Style(root)
    style.theme_use("clam")

    # Общий фон
    style.configure(".", background=CLR_BG, foreground=CLR_TEXT, font=("Segoe UI", 10))

    # Frame / LabelFrame
    style.configure("TFrame", background=CLR_BG)
    style.configure("Card.TFrame", background=CLR_PANEL, relief="flat")
    style.configure("TLabelframe", background=CLR_PANEL, relief="flat",
                    bordercolor=CLR_BORDER, lightcolor=CLR_BORDER, darkcolor=CLR_BORDER)
    style.configure("TLabelframe.Label", background=CLR_PANEL,
                    foreground=CLR_ACCENT, font=("Segoe UI", 10, "bold"))

    # Кнопки — основная
    style.configure("Accent.TButton",
                    background=CLR_ACCENT, foreground="#ffffff",
                    font=("Segoe UI", 10, "bold"),
                    padding=(14, 8), relief="flat", borderwidth=0)
    style.map("Accent.TButton",
              background=[("active", CLR_ACCENT_H), ("pressed", CLR_ACCENT_H)],
              foreground=[("active", "#ffffff")])

    # Кнопки — опасная (удалить)
    style.configure("Danger.TButton",
                    background=CLR_DANGER, foreground="#ffffff",
                    font=("Segoe UI", 9), padding=(10, 6), relief="flat", borderwidth=0)
    style.map("Danger.TButton",
              background=[("active", CLR_DANGER_H), ("pressed", CLR_DANGER_H)],
              foreground=[("active", "#ffffff")])

    # Кнопки — вторичная
    style.configure("TButton",
                    background="#e8eaf0", foreground=CLR_TEXT,
                    font=("Segoe UI", 10), padding=(12, 7), relief="flat", borderwidth=0)
    style.map("TButton",
              background=[("active", "#d0d4e0"), ("pressed", "#c5c9d8")],
              foreground=[("active", CLR_TEXT)])

    # Entry
    style.configure("TEntry", fieldbackground=CLR_ENTRY_BG, foreground=CLR_TEXT,
                    bordercolor=CLR_BORDER, lightcolor=CLR_BORDER, darkcolor=CLR_BORDER,
                    insertcolor=CLR_ACCENT, padding=6)
    style.map("TEntry", bordercolor=[("focus", CLR_ACCENT)])

    # Label
    style.configure("TLabel", background=CLR_PANEL, foreground=CLR_TEXT, font=("Segoe UI", 10))
    style.configure("Title.TLabel", background=CLR_BG, foreground=CLR_TEXT,
                    font=("Segoe UI", 15, "bold"))
    style.configure("Sub.TLabel", background=CLR_PANEL, foreground=CLR_SUBTEXT,
                    font=("Segoe UI", 9))
    style.configure("FieldLabel.TLabel", background=CLR_PANEL, foreground=CLR_SUBTEXT,
                    font=("Segoe UI", 9, "bold"))

    # Radiobutton
    style.configure("TRadiobutton", background=CLR_PANEL, foreground=CLR_TEXT,
                    font=("Segoe UI", 10))

    # Scrollbar
    style.configure("Vertical.TScrollbar", background=CLR_BG, troughcolor=CLR_BG,
                    bordercolor=CLR_BG, arrowcolor=CLR_SUBTEXT, relief="flat")

    # PanedWindow
    style.configure("TPanedwindow", background=CLR_BG)

    fio_var = tk.StringVar()
    position_var = tk.StringVar()
    subject_var = tk.StringVar()
    gender_var = tk.StringVar(value="male")
    recipient_name_var = tk.StringVar()
    recipient_position_var = tk.StringVar()
    recipient_organization_var = tk.StringVar()
    document_date_var = tk.StringVar(value=datetime.datetime.now().strftime("%d.%m.%Y"))

    app_frames = []
    current_pdf_path = None

    try:
        if os.name == 'nt':
            font_paths = [
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/times.ttf",
                "C:/Windows/Fonts/calibri.ttf"
            ]

        font_registered = False
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('RussianFont', font_path))
                addMapping('RussianFont', 0, 0, 'RussianFont')  # normal
                addMapping('RussianFont', 1, 0, 'RussianFont')  # bold
                addMapping('RussianFont', 0, 1, 'RussianFont')  # italic
                addMapping('RussianFont', 1, 1, 'RussianFont')  # bold italic
                font_registered = True
                print(f"Зарегистрирован шрифт: {font_path}")
                break

        if not font_registered:
            pdfmetrics.registerFont(TTFont('DejaVu', 'DejaVuSans.ttf'))
            addMapping('DejaVu', 0, 0, 'DejaVu')
            print("Используется шрифт DejaVuSans")

    except Exception as e:
        print(f"Ошибка регистрации шрифта: {e}")
        pass

    def get_gender_suffixes(gender):
        if gender == "male":
            return {
                "occupies": "занимающий",
                "prepared": "подготовил",
                "reports": "прошу рассмотреть"
            }
        else:
            return {
                "occupies": "занимающая",
                "prepared": "подготовила",
                "reports": "прошу рассмотреть"
            }

    def validate_date(date_string):
        """Validate date format DD.MM.YYYY"""
        try:
            parts = date_string.strip().split('.')
            if len(parts) != 3:
                return False, "Неверный формат даты"
            
            day, month, year = parts
            
            if not (day.isdigit() and month.isdigit() and year.isdigit()):
                return False, "Дата должна содержать только цифры"
            
            day, month, year = int(day), int(month), int(year)
            
            if year < 1900 or year > 2100:
                return False, "Год должен быть между 1900 и 2100"
            
            # Try to create a datetime object to validate the date
            datetime.datetime(year, month, day)
            return True, "OK"
        except ValueError:
            return False, "Некорректная дата"

    def on_date_change(*args):
        """Handle date input changes with validation"""
        current_text = document_date_var.get()
        
        # Allow empty input during typing
        if not current_text:
            date_status_label.config(text="", foreground="#6b7280")
            return
        
        is_valid, message = validate_date(current_text)
        
        if is_valid:
            date_status_label.config(text="✓ Дата корректна", foreground="#10b981")
            update_preview()
        else:
            date_status_label.config(text=f"✗ {message}", foreground="#e84545")

    def validate_form():
        errors = []

        if not fio_var.get().strip():
            errors.append("ФИО отправителя")
        if not position_var.get().strip():
            errors.append("Должность отправителя")
        if not subject_var.get().strip():
            errors.append("Тема заявки")
        if not recipient_name_var.get().strip():
            errors.append("ФИО получателя")
        if not recipient_position_var.get().strip():
            errors.append("Должность получателя")
        
        # Validate date
        date_text = document_date_var.get().strip()
        if not date_text:
            errors.append("Дата документа")
        else:
            is_valid, message = validate_date(date_text)
            if not is_valid:
                errors.append(f"Дата документа ({message})")

        if errors:
            messagebox.showwarning("Ошибка валидации",
                                   f"Пожалуйста, заполните следующие обязательные поля:\n" + "\n".join(
                                       f"• {err}" for err in errors))
            return False
        return True

    def delete_app(frame):
        if len(app_frames) > 1:
            frame.destroy()
            app_frames.remove(frame)
            renumber_apps()
            update_preview()
        else:
            messagebox.showwarning("Предупреждение", "Должно быть хотя бы одно приложение")

    def renumber_apps():
        for idx, app_frame in enumerate(app_frames, 1):
            app_frame.configure(text=f"Приложение {idx}")

    def create_app_frame(app_num, title="", text=""):
        frame = ttk.LabelFrame(apps_container, text=f"Приложение {app_num}", padding="10")
        frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(frame, text="Заголовок:", font=('Arial', 9, 'bold')).pack(anchor=tk.W, pady=(0, 5))
        title_entry = ttk.Entry(frame, width=50, font=('Arial', 9))
        title_entry.pack(fill=tk.X, pady=(0, 10))
        title_entry.insert(0, title)
        title_entry.bind('<KeyRelease>', lambda e: update_preview())

        ttk.Label(frame, text="Текст приложения:", font=('Arial', 9, 'bold')).pack(anchor=tk.W, pady=(0, 5))

        text_frame = ttk.Frame(frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 5))

        text_widget = tk.Text(text_frame, height=6, width=60, font=('Arial', 9), wrap=tk.WORD)
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_widget.insert("1.0", text)
        text_widget.bind('<KeyRelease>', lambda e: update_preview())

        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        text_widget.configure(yscrollcommand=scrollbar.set)

        delete_btn = ttk.Button(frame, text="Удалить приложение",
                                command=lambda: delete_app(frame))
        delete_btn.pack(pady=(5, 0))

        frame.title_entry = title_entry
        frame.text_widget = text_widget

        return frame

    def add_application():
        app_num = len(app_frames) + 1
        app_frame = create_app_frame(app_num)
        app_frames.append(app_frame)
        update_preview()

    def get_applications_data():
        apps_data = []
        for app_frame in app_frames:
            title = app_frame.title_entry.get().strip()
            text = app_frame.text_widget.get("1.0", tk.END).strip()

            if not text:
                continue

            if not title:
                title = "Без названия"

            apps_data.append({
                "title": title,
                "text": text
            })

        return apps_data

    def create_pdf_preview():
        nonlocal current_pdf_path

        if current_pdf_path and os.path.exists(current_pdf_path):
            try:
                os.unlink(current_pdf_path)
            except:
                pass

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        current_pdf_path = temp_file.name
        temp_file.close()

        doc = SimpleDocTemplate(current_pdf_path, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Normal'],
            fontName='RussianFont',
            fontSize=16,
            alignment=1,
            spaceAfter=30,
            textColor=colors.HexColor('#1a1a1a'),
            leading=20
        )

        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontName='RussianFont',
            fontSize=12,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12,
            spaceBefore=12,
            leading=14
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='RussianFont',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#444444')
        )

        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontName='RussianFont',
            fontSize=10,
            alignment=2,  # Право
            textColor=colors.HexColor('#666666'),
            leading=12
        )

        app_title_style = ParagraphStyle(
            'AppTitle',
            parent=styles['Normal'],
            fontName='RussianFont',
            fontSize=12,
            alignment=2,  # Право
            textColor=colors.HexColor('#1a1a1a'),
            leading=14
        )

        fio = fio_var.get().strip()
        position = position_var.get().strip()
        subject = subject_var.get().strip()
        gender = gender_var.get()
        recipient_name = recipient_name_var.get().strip()
        recipient_position = recipient_position_var.get().strip()
        recipient_org = recipient_organization_var.get().strip()
        document_date = document_date_var.get().strip()

        suffixes = get_gender_suffixes(gender)

        story = []

        story.append(Paragraph("ЗАЯВКА", title_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph(f"Дата: {document_date}", date_style))
        story.append(Spacer(1, 20))

        story.append(Paragraph("КОМУ:", header_style))
        recipient_info = []
        if recipient_name:
            recipient_info.append(f"<b>ФИО:</b> {recipient_name}")
        if recipient_position:
            recipient_info.append(f"<b>Должность:</b> {recipient_position}")
        if recipient_org:
            recipient_info.append(f"<b>Организация:</b> {recipient_org}")

        for info in recipient_info:
            story.append(Paragraph(info, normal_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("ОТ КОГО:", header_style))
        sender_info = []
        if fio:
            sender_info.append(f"<b>ФИО:</b> {fio}")
        if position:
            sender_info.append(f"<b>Должность:</b> {position}")

        for info in sender_info:
            story.append(Paragraph(info, normal_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("ТЕМА:", header_style))
        story.append(Paragraph(f"<b>{subject}</b>", normal_style))
        story.append(Spacer(1, 20))

        if fio and position and subject:
            content = f"""Уважаемый(ая) {recipient_name if recipient_name else 'руководитель'}!<br/><br/>
{suffixes['reports']}, нижеуказанную заявку. Я, {fio}, {suffixes['occupies']} должность {position}, 
{suffixes['prepared']} данную заявку по вопросу: "{subject}".<br/><br/>"""

            story.append(Paragraph(content, normal_style))

            apps_data = get_applications_data()
            if apps_data:
                story.append(Paragraph("К настоящей заявке прилагаются следующие документы:", normal_style))
                for idx, app in enumerate(apps_data, 1):
                    if len(apps_data) == 1:
                        story.append(Paragraph(f"• Приложение: {app['title']}", normal_style))
                    else:
                        story.append(Paragraph(f"• Приложение {idx}: {app['title']}", normal_style))
                story.append(Spacer(1, 10))

            content_end = f"""Прошу Вас рассмотреть данную заявку и принять соответствующее решение.<br/><br/>
Дата составления: {document_date}<br/><br/>
С уважением,<br/>
{fio}<br/>
{position}<br/><br/>
___________________<br/>
(подпись)"""

            story.append(Paragraph(content_end, normal_style))

            if apps_data:
                story.append(PageBreak())
                story.append(Paragraph("ПРИЛОЖЕНИЯ", title_style))
                story.append(Spacer(1, 20))

                for idx, app in enumerate(apps_data, 1):
                    if len(apps_data) == 1:
                        app_title = "Приложение"
                    else:
                        app_title = f"Приложение {idx}"

                    story.append(Paragraph(app_title, app_title_style))
                    story.append(Spacer(1, 10))

                    story.append(Paragraph(f"<b>{app['title']}</b>", header_style))
                    story.append(Spacer(1, 10))

                    app_text = app['text'].replace('\n', '<br/>')
                    story.append(Paragraph(app_text, normal_style))
                    story.append(Spacer(1, 20))

                    if idx < len(apps_data):
                        story.append(Spacer(1, 20))

        doc.build(story)
        return current_pdf_path

    def display_pdf_in_frame(pdf_path, frame_widget):
        for widget in frame_widget.winfo_children():
            widget.destroy()

        if not pdf_path or not os.path.exists(pdf_path):
            error_label = tk.Label(frame_widget,
                                   text="PDF файл не создан\nПожалуйста, заполните обязательные поля",
                                   font=('Arial', 12),
                                   fg='orange',
                                   bg='white')
            error_label.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)
            return

        try:
            doc = fitz.open(pdf_path)

            canvas_frame = tk.Canvas(frame_widget, bg='white')
            canvas_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            scrollbar = ttk.Scrollbar(frame_widget, orient="vertical", command=canvas_frame.yview)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            canvas_frame.configure(yscrollcommand=scrollbar.set)

            inner_frame = tk.Frame(canvas_frame, bg='white')
            canvas_frame.create_window((0, 0), window=inner_frame, anchor="nw")

            images = []
            for page_num in range(len(doc)):
                page = doc[page_num]

                rect = page.rect

                target_width = 550
                zoom = target_width / rect.width
                mat = fitz.Matrix(zoom, zoom)

                pix = page.get_pixmap(matrix=mat, alpha=False)

                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                photo = ImageTk.PhotoImage(img)
                images.append(photo)

                page_label = tk.Label(inner_frame, image=photo, bg='white')
                page_label.pack(pady=10)

                page_num_label = tk.Label(inner_frame, text=f"Страница {page_num + 1}",
                                          font=('Arial', 9), fg='gray', bg='white')
                page_num_label.pack(pady=(0, 20))

            inner_frame.update_idletasks()
            canvas_frame.configure(scrollregion=canvas_frame.bbox("all"))

            canvas_frame.images = images

            def on_mousewheel(event):
                canvas_frame.yview_scroll(int(-1 * (event.delta / 120)), "units")

            canvas_frame.bind("<MouseWheel>", on_mousewheel)

            doc.close()

        except Exception as e:
            error_label = tk.Label(frame_widget,
                                   text=f"Ошибка отображения PDF:\n{str(e)}",
                                   font=('Arial', 10),
                                   fg='red',
                                   bg='white',
                                   wraplength=500)
            error_label.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)

    def update_preview():
        try:
            if not fio_var.get().strip() or not position_var.get().strip() or not subject_var.get().strip():
                for widget in pdf_preview_frame.winfo_children():
                    widget.destroy()
                info_label = tk.Label(pdf_preview_frame,
                                      text="Заполните обязательные поля\n(ФИО, должность отправителя и тему заявки)\nдля отображения предпросмотра",
                                      font=('Arial', 12),
                                      fg='orange',
                                      bg='white',
                                      wraplength=500)
                info_label.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)
                return

            # Check if date is valid
            date_text = document_date_var.get().strip()
            if date_text:
                is_valid, _ = validate_date(date_text)
                if not is_valid:
                    for widget in pdf_preview_frame.winfo_children():
                        widget.destroy()
                    info_label = tk.Label(pdf_preview_frame,
                                          text="Проверьте формат даты\n(ДД.МММ.ГГГГ)",
                                          font=('Arial', 12),
                                          fg='orange',
                                          bg='white',
                                          wraplength=500)
                    info_label.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)
                    return

            pdf_path = create_pdf_preview()

            display_pdf_in_frame(pdf_path, pdf_preview_frame)

        except Exception as e:
            for widget in pdf_preview_frame.winfo_children():
                widget.destroy()

            error_label = tk.Label(pdf_preview_frame,
                                   text=f"Ошибка создания PDF:\n{str(e)}\n\nПроверьте заполнение всех полей",
                                   font=('Arial', 10),
                                   fg='red',
                                   bg='white',
                                   wraplength=500)
            error_label.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)

    def generate():
        if not validate_form():
            return

        fio = fio_var.get().strip()
        position = position_var.get().strip()
        subject = subject_var.get().strip()
        gender = gender_var.get()
        recipient_name = recipient_name_var.get().strip()
        recipient_position = recipient_position_var.get().strip()
        recipient_org = recipient_organization_var.get().strip()
        document_date = document_date_var.get().strip()

        apps_data = get_applications_data()

        try:
            doc = docx.Document()

            title = doc.add_heading('ЗАЯВКА', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.add_run(f"Дата: {document_date}").bold = True

            doc.add_paragraph()

            doc.add_heading('Получатель:', 2)
            if recipient_name:
                doc.add_paragraph(f"ФИО: {recipient_name}")
            if recipient_position:
                doc.add_paragraph(f"Должность: {recipient_position}")
            if recipient_org:
                doc.add_paragraph(f"Организация: {recipient_org}")

            doc.add_paragraph()

            doc.add_heading('Отправитель:', 2)
            doc.add_paragraph(f"ФИО: {fio}\nДолжность: {position}")

            doc.add_paragraph()

            doc.add_heading('Тема:', 2)
            doc.add_paragraph(subject).bold = True

            doc.add_paragraph()

            suffixes = get_gender_suffixes(gender)
            doc.add_heading('Содержание:', 2)

            content = f"""Уважаемый(ая) {recipient_name if recipient_name else 'руководитель'}!

{suffixes['reports']}, нижеуказанную заявку. Я, {fio}, {suffixes['occupies']} должность {position}, 
{suffixes['prepared']} данную заявку по вопросу: "{subject}".

"""
            doc.add_paragraph(content)

            if apps_data:
                doc.add_paragraph("К настоящей заявке прилагаются следующие документы:")

                apps_list_para = doc.add_paragraph()
                for idx, app in enumerate(apps_data, 1):
                    if len(apps_data) == 1:
                        apps_list_para.add_run(f"Приложение: {app['title']}")
                    else:
                        apps_list_para.add_run(f"Приложение {idx}: {app['title']}")

                    if idx < len(apps_data):
                        apps_list_para.add_run("\n")

                doc.add_paragraph()

            doc.add_paragraph(f"""Прошу Вас рассмотреть данную заявку и принять соответствующее решение.

Дата составления: {document_date}

С уважением,
{fio},
{position}""")

            doc.add_paragraph()

            signature_para = doc.add_paragraph()
            signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            signature_para.add_run("___________________").bold = True
            signature_para.add_run("\n(подпись)")

            if apps_data:
                doc.add_page_break()
                doc.add_heading('ПРИЛОЖЕНИЯ', 1)
                doc.add_paragraph()

                for idx, app in enumerate(apps_data, 1):
                    if len(apps_data) == 1:
                        app_title_para = doc.add_paragraph()
                        app_title_para.add_run("Приложение").bold = True
                        app_title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        app_title_para = doc.add_paragraph()
                        app_title_para.add_run(f"Приложение {idx}").bold = True
                        app_title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    doc.add_paragraph()

                    title_para = doc.add_paragraph()
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.add_run(app['title'])
                    title_run.bold = True
                    title_run.font.size = Pt(14)

                    doc.add_paragraph()

                    doc.add_paragraph(app['text'])

                    if idx < len(apps_data):
                        doc.add_page_break()

            filename = f"заявка_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)

            pdf_filename = f"заявка_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            if current_pdf_path and os.path.exists(current_pdf_path):
                import shutil
                shutil.copy(current_pdf_path, pdf_filename)

            messagebox.showinfo("Успех",
                                f"Документы успешно созданы!\n"
                                f"Сохранены как:\n"
                                f"• {filename}\n"
                                f"• {pdf_filename}\n\n"
                                f"Создано приложений: {len(apps_data)}")
            os.startfile(os.path.dirname(os.path.abspath(filename)))

        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка:\n{str(e)}")

    main_paned = ttk.PanedWindow(root, orient=tk.HORIZONTAL)
    main_paned.pack(fill=tk.BOTH, expand=True)

    left_frame = ttk.Frame(main_paned)
    main_paned.add(left_frame, weight=1)

    left_canvas = tk.Canvas(left_frame)
    left_scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=left_canvas.yview)
    scrollable_frame = ttk.Frame(left_canvas)

    scrollable_frame.bind("<Configure>", lambda e: left_canvas.configure(scrollregion=left_canvas.bbox("all")))
    left_canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    left_canvas.configure(yscrollcommand=left_scrollbar.set)

    left_canvas.pack(side="left", fill="both", expand=True)
    left_scrollbar.pack(side="right", fill="y")

    def on_mousewheel_left(event):
        left_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    left_canvas.bind("<MouseWheel>", on_mousewheel_left)

    ttk.Label(scrollable_frame, text="Генератор документов с приложениями", font=('Arial', 14, 'bold')).pack(
        pady=(0, 20))

    recipient_frame = ttk.LabelFrame(scrollable_frame, text="Информация о получателе", padding="10")
    recipient_frame.pack(fill=tk.X, pady=(0, 15))

    ttk.Label(recipient_frame, text="ФИО получателя:*", font=('Arial', 10)).grid(row=0, column=0, sticky=tk.W,
                                                                                 pady=(0, 5), padx=(0, 10))
    ttk.Entry(recipient_frame, textvariable=recipient_name_var, width=40, font=('Arial', 10)).grid(row=0, column=1,
                                                                                                   pady=(0, 5))
    recipient_name_var.trace('w', lambda *args: update_preview())

    ttk.Label(recipient_frame, text="Должность получателя:*", font=('Arial', 10)).grid(row=1, column=0, sticky=tk.W,
                                                                                       pady=(0, 5), padx=(0, 10))
    ttk.Entry(recipient_frame, textvariable=recipient_position_var, width=40, font=('Arial', 10)).grid(row=1, column=1,
                                                                                                       pady=(0, 5))
    recipient_position_var.trace('w', lambda *args: update_preview())

    ttk.Label(recipient_frame, text="Организация получателя:", font=('Arial', 10)).grid(row=2, column=0, sticky=tk.W,
                                                                                        pady=(0, 5), padx=(0, 10))
    ttk.Entry(recipient_frame, textvariable=recipient_organization_var, width=40, font=('Arial', 10)).grid(row=2,
                                                                                                           column=1,
                                                                                                           pady=(0, 5))
    recipient_organization_var.trace('w', lambda *args: update_preview())

    sender_frame = ttk.LabelFrame(scrollable_frame, text="Информация об отправителе", padding="10")
    sender_frame.pack(fill=tk.X, pady=(0, 15))

    ttk.Label(sender_frame, text="ФИО отправителя:*", font=('Arial', 10)).grid(row=0, column=0, sticky=tk.W,
                                                                               pady=(0, 5), padx=(0, 10))
    ttk.Entry(sender_frame, textvariable=fio_var, width=40, font=('Arial', 10)).grid(row=0, column=1, pady=(0, 5))
    fio_var.trace('w', lambda *args: update_preview())

    ttk.Label(sender_frame, text="Должность отправителя:*", font=('Arial', 10)).grid(row=1, column=0, sticky=tk.W,
                                                                                     pady=(0, 5), padx=(0, 10))
    ttk.Entry(sender_frame, textvariable=position_var, width=40, font=('Arial', 10)).grid(row=1, column=1, pady=(0, 5))
    position_var.trace('w', lambda *args: update_preview())

    gender_frame = ttk.LabelFrame(sender_frame, text="Пол отправителя", padding="5")
    gender_frame.grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(10, 0))

    ttk.Radiobutton(gender_frame, text="Мужской", variable=gender_var, value="male", command=update_preview).pack(
        side=tk.LEFT, padx=(0, 20))
    ttk.Radiobutton(gender_frame, text="Женский", variable=gender_var, value="female", command=update_preview).pack(
        side=tk.LEFT)

    letter_frame = ttk.LabelFrame(scrollable_frame, text="Основная информация о заявке", padding="10")
    letter_frame.pack(fill=tk.X, pady=(0, 15))

    ttk.Label(letter_frame, text="Тема заявки:*", font=('Arial', 10)).grid(row=0, column=0, sticky=tk.W, pady=(0, 5),
                                                                           padx=(0, 10))
    ttk.Entry(letter_frame, textvariable=subject_var, width=40, font=('Arial', 10)).grid(row=0, column=1, pady=(0, 5))
    subject_var.trace('w', lambda *args: update_preview())

    ttk.Label(letter_frame, text="Дата документа:*", font=('Arial', 10)).grid(row=1, column=0, sticky=tk.W, pady=(0, 5),
                                                                              padx=(0, 10))
    date_entry = ttk.Entry(letter_frame, textvariable=document_date_var, width=40, font=('Arial', 10))
    date_entry.grid(row=1, column=1, pady=(0, 5))
    document_date_var.trace('w', on_date_change)

    date_status_label = ttk.Label(letter_frame, text="✓ Дата корректна", font=('Arial', 9), foreground='#10b981')
    date_status_label.grid(row=2, column=1, sticky=tk.W, pady=(2, 10))

    hint_label = ttk.Label(letter_frame, text="Формат: ДД.МММ.ГГГГ (например: 29.05.2026)", font=('Arial', 8), foreground='#6b7280')
    hint_label.grid(row=3, column=0, columnspan=2, sticky=tk.W, pady=(0, 5))

    apps_container = ttk.LabelFrame(scrollable_frame, text="Приложения к заявке", padding="10")
    apps_container.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

    add_app_btn = ttk.Button(apps_container, text="+ Добавить приложение", command=add_application)
    add_app_btn.pack(pady=(0, 10))

    default_app = create_app_frame(1)
    app_frames.append(default_app)

    generate_btn = ttk.Button(scrollable_frame, text="Сгенерировать документ", command=generate)
    generate_btn.pack(pady=10)

    info_label = ttk.Label(scrollable_frame, text="* - обязательные поля\n"
                                                  "Нажмите кнопку для создания документа Word и PDF\n"
                                                  "Документы будут сохранены в текущей папке\n\n"
                                                  "Примечание: приложения без текста не будут добавлены в документ",
                           font=('Arial', 9), foreground='gray')
    info_label.pack(pady=(15, 0))

    right_frame = ttk.Frame(main_paned)
    main_paned.add(right_frame, weight=1)

    preview_label = ttk.Label(right_frame, text="Предпросмотр документа (PDF)", font=('Arial', 12, 'bold'))
    preview_label.pack(pady=(10, 5))

    refresh_btn = ttk.Button(right_frame, text="Обновить предпросмотр", command=update_preview)
    refresh_btn.pack(pady=5)

    pdf_preview_frame = ttk.Frame(right_frame, relief=tk.SUNKEN, borderwidth=1)
    pdf_preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

    update_preview()

    def on_closing():
        nonlocal current_pdf_path
        if current_pdf_path and os.path.exists(current_pdf_path):
            try:
                os.unlink(current_pdf_path)
            except:
                pass
        root.destroy()

    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()

if __name__ == "__main__":
    main()
